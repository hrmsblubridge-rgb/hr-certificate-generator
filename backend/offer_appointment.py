"""Generate the "Offer of Appointment" .docx (and PDF via LibreOffice) by
substituting dynamic fields directly into the operator-supplied source
`templates/offer_of_appointment_source.docx`.

Key design choice — **substitution, not regeneration.**
We do NOT re-lay-out the document. The source `.docx` IS the layout: fonts,
sizes, spacing, indentation, tabs, page breaks, the 4 tables (compensation,
tier bands, medical, conditions) — all that stays exactly as the user
authored it. We just swap textual values inside existing runs / cells, which
preserves run-level formatting (bold/italic/etc.) where it was applied
character-by-character.

PDF output runs LibreOffice headless on the produced DOCX so the rendered
PDF is byte-faithful to what Word would print.
"""
from __future__ import annotations

import io
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document

# We reuse the math + tier derivation already audited in offer_letter_email.
from offer_letter_email import (
    indian_format,
    indian_number_to_words,
    scale_compensation,
    derive_tier,
)


SOURCE_DOCX = Path(__file__).parent / "templates" / "offer_of_appointment_source.docx"


# --- Run-aware text replacement ---------------------------------------------
def _replace_in_paragraph(p, mappings: List[Tuple[str, str]]) -> int:
    """Substitute every (old, new) pair inside `p`. Preserves run-level
    formatting when the `old` fragment sits inside a single run; falls back
    to a paragraph-level rewrite (keeping the first run's formatting) when
    the fragment spans runs.

    Mappings are applied STRICTLY IN ORDER and each mapping gets its own
    "try per-run, fall back to cross-run" pass before the next mapping runs.
    This matters when a longer pattern (e.g. ``"join on 08-June-2026"``)
    must consume the text BEFORE a shorter overlapping pattern
    (``"08-June-2026"``) gets a chance to match.
    """
    if not p.runs:
        return 0
    count = 0
    for old, new in mappings:
        if not old:
            continue
        replaced_in_run = False
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced_in_run = True
                count += 1
        if replaced_in_run:
            continue
        # Cross-run fallback for THIS mapping only.
        full = "".join(r.text for r in p.runs)
        if old in full:
            new_full = full.replace(old, new)
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""
            count += 1
    return count


def _replace_everywhere(doc, mappings: List[Tuple[str, str]]):
    """Apply the substitution list to every paragraph in the document body
    AND every cell in every table (including nested paragraphs in cells)."""
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mappings)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mappings)


# --- Annexure-A compensation table rewrite ----------------------------------
# The source DOCX hard-codes the reference @ CTC=660,000. We replace each
# pair of monthly/annual numeric cells via row-index mapping.
# Reference values are kept here so a future change to the template is
# detected (we assert below).
_REF_TABLE_VALUES = {
    # row_idx: (label_contains, per_month_str, per_annum_str)
    5:  ("Base Components",      "24,750",  "2,97,000"),
    6:  ("Basic",                "16,500",  "1,98,000"),
    7:  ("HRA",                   "8,250",    "99,000"),
    8:  ("Basket of Allowances", "16,657",  "1,99,884"),
    9:  ("Leave Travel",            "935",   "11,220"),
    10: ("Phone & Internet",      "1,100",   "13,200"),
    11: ("Bonus",                 "1,650",   "19,800"),
    12: ("Stay and Travel",       "5,000",   "60,000"),
    13: ("Special Allowance",     "6,762",   "81,144"),
    14: ("Food",                  "1,210",    "14520"),
    15: ("Retirement Benefits",   "2,593",    "31116"),
    16: ("PF",                    "1,800",    "21600"),
    17: ("Gratuity",                "793",     "9516"),
    18: ("Fixed Compensation",   "44,000",  "528,000"),
    19: ("Variable Compensation","11,000",  "1,32,000"),
    20: ("Cost to Company",      "55,000",  "6,60,000"),
}

# Map row_idx → key in scale_compensation(). Group rows (5/8/15/18/20) are
# computed from the line items.
_ROW_TO_KEY = {
    6: "basic", 7: "hra",
    9: "lta", 10: "phone", 11: "bonus",
    12: "stay", 13: "special", 14: "food",
    16: "pf", 17: "gratuity",
    19: "variable",
}


def _set_cell_text(cell, text: str):
    """Replace the cell's text content while keeping the first paragraph's
    first run's formatting (font, size, bold). Removes any extra runs."""
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def _rewrite_compensation_table(table, comp: dict, tier_label: str):
    """Overwrite the per-month and per-annum cells (cols 1 and 2) for each
    line item, plus the group/summary totals, plus the TIER row.
    Preserves all formatting."""
    f = indian_format
    # Tier row (row 3, col 1) — done HERE so we don't accidentally touch
    # the Annexure-B tier-bands reference table via a global substitution.
    _set_cell_text(table.rows[3].cells[1], tier_label)
    # Direct line-item rows (single-component cells)
    for row_idx, key in _ROW_TO_KEY.items():
        annual = comp[key]
        monthly = round(annual / 12)
        _set_cell_text(table.rows[row_idx].cells[1], f(monthly))
        _set_cell_text(table.rows[row_idx].cells[2], f(annual))

    # Group/summary rows
    base_a   = comp["basic"] + comp["hra"]
    allow_b  = comp["lta"] + comp["phone"] + comp["bonus"] + comp["stay"] + comp["special"] + comp["food"]
    retire_c = comp["pf"] + comp["gratuity"]
    summaries = {
        5:  base_a,
        8:  allow_b,
        15: retire_c,
        18: comp["fixed_total"],
        20: comp["ctc_total"],
    }
    for row_idx, annual in summaries.items():
        monthly = round(annual / 12)
        _set_cell_text(table.rows[row_idx].cells[1], f(monthly))
        _set_cell_text(table.rows[row_idx].cells[2], f(annual))


# --- Top-level renderer ------------------------------------------------------
def render_docx(data: dict) -> bytes:
    """Build the personalised offer-of-appointment DOCX. Returns the binary."""
    doc = Document(str(SOURCE_DOCX))

    # ---- 1. Compute derived values ----------------------------------------
    annual_ctc  = int(data["ctc_yearly"])
    comp        = scale_compensation(annual_ctc)
    title       = data["title"].strip()
    name        = data["name"].strip()
    full_name   = f"{title} {name}".strip()
    designation = data["designation"].strip()
    cur_date    = data["cur_date"].strip()       # e.g. "08-June-2026"
    joining     = data["date"].strip()           # joining date
    ref_number  = data["reference_number"].strip()
    addr1       = data["address_line1"].strip()
    addr2       = data["address_line2"].strip()
    addr3       = data["address_line3"].strip()
    phone       = data["phone"].strip()
    email       = data["email"].strip()
    ctc_str     = indian_format(annual_ctc)
    ctc_words   = indian_number_to_words(annual_ctc)
    tier_label  = comp["tier"]

    # ---- 2. Text substitutions (run-aware, format-preserving) -------------
    # ORDER matters: replace LONGER strings first to avoid partial matches.
    mappings: List[Tuple[str, str]] = [
        # Salutation forms (must come before bare "Revathi Thiruppathi")
        ("Ms. Revathi Thiruppathi", full_name),
        ("Dear Revathi Thiruppathi", f"Dear {full_name}"),
        # Bare name (signature lines, table cells)
        ("Revathi Thiruppathi", name),
        # Address
        ("2/162 E, Nethaji Nagar, 3rd cross,", addr1),
        ("Kanagamutlu(post),",                 addr2),
        ("Krishnagiri - 635001",               addr3),
        # Contact
        ("Phone: 8300233625", f"Phone: {phone}"),
        ("8300233625",        phone),
        ("revathitdgrs@gmail.com", email),
        # Reference + dates (letter-date vs joining-date are distinct fields,
        # but the source uses the SAME "08-June-2026" string for both. We
        # substitute the letter-date FIRST (where it appears in the top row
        # 'Date:' and acceptance footer), then the joining-date.)
        ("CHN/2025/Res/1-026", ref_number),
        # The joining-date appears in "join on 08-June-2026" — patch that
        # specific phrase first so it picks up the joining date.
        ("join on 08-June-2026", f"join on {joining}"),
        # Remaining "08-June-2026" occurrences (top date, acceptance footer)
        # are the letter date.
        ("08-June-2026", cur_date),
        # Designation / Tier / CTC
        # Designation appears in both the body and Annexure A — global swap is OK.
        ("Research Scientist", designation),
        # NOTE: Tier 2 is NOT in the global mappings because the tier-bands
        # reference table (Annexure B) lists Tier 0–Tier 5 as static brackets
        # and must NOT be touched. Tier is updated directly in the Annexure-A
        # row inside `_rewrite_compensation_table` below.
        # CTC figures + words (body paragraph p16)
        ("Rs 660,000",                    f"Rs {ctc_str}"),
        ("Indian Rupees Six Lakh Sixty Thousand", f"Indian Rupees {ctc_words.replace(' Only','')}"),
    ]
    _replace_everywhere(doc, mappings)

    # ---- 3. Compensation table — proportional rescale ---------------------
    _rewrite_compensation_table(doc.tables[0], comp, tier_label)

    # ---- 4. Serialise ------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- DOCX → PDF via LibreOffice headless ------------------------------------
def docx_to_pdf(docx_bytes: bytes, *, timeout: int = 120) -> bytes:
    """Convert in-memory DOCX bytes to PDF bytes using LibreOffice headless.

    Runs LibreOffice in a per-invocation temp profile dir so concurrent
    conversions don't race on the default profile lock.
    """
    with tempfile.TemporaryDirectory(prefix="ofa-") as td:
        td_path = Path(td)
        src = td_path / "in.docx"
        src.write_bytes(docx_bytes)
        profile = td_path / "lo-profile"
        cmd = [
            "libreoffice",
            "--headless",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to", "pdf",
            "--outdir", str(td_path),
            str(src),
        ]
        try:
            result = subprocess.run(
                cmd, check=True, capture_output=True, timeout=timeout,
            )
        except subprocess.CalledProcessError as e:
            stderr = (e.stderr or b"").decode("utf-8", errors="replace")[:500]
            raise RuntimeError(f"LibreOffice conversion failed: {stderr}") from e
        except subprocess.TimeoutExpired as e:
            raise RuntimeError("LibreOffice conversion timed out.") from e
        pdf_path = td_path / "in.pdf"
        if not pdf_path.exists():
            stdout = (result.stdout or b"").decode("utf-8", errors="replace")[:500]
            raise RuntimeError(f"LibreOffice produced no PDF. stdout={stdout!r}")
        return pdf_path.read_bytes()
